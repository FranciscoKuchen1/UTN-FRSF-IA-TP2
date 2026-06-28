from pathlib import Path

import fitz
import pytest
from docx import Document

from rag.ingest import chunk_text, parse_file


def test_parse_txt_supports_utf8_and_cp1252(tmp_path):
    utf8_file = tmp_path / "utf8.txt"
    utf8_file.write_text("Declaracion de IVA", encoding="utf-8")
    cp1252_file = tmp_path / "windows.txt"
    cp1252_file.write_bytes("Declaración de IVA".encode("cp1252"))

    assert "IVA" in parse_file(utf8_file)
    assert "Declaración" in parse_file(cp1252_file)


def test_parse_docx_includes_paragraphs_and_tables(tmp_path):
    path = tmp_path / "guia.docx"
    document = Document()
    document.add_paragraph("Guia tributaria")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "IVA"
    table.cell(0, 1).text = "18/06/2026"
    document.save(path)

    text = parse_file(path)

    assert "Guia tributaria" in text
    assert "IVA | 18/06/2026" in text


def test_parse_pdf_extracts_text(tmp_path):
    path = tmp_path / "calendario.pdf"
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "Calendario IVA 2026")
    document.save(path)
    document.close()

    assert "Calendario IVA 2026" in parse_file(path)


def test_parse_file_rejects_unsupported_or_empty_documents(tmp_path):
    unsupported = tmp_path / "planilla.xlsx"
    unsupported.write_bytes(b"data")
    empty = tmp_path / "vacio.txt"
    empty.write_text("   ", encoding="utf-8")

    with pytest.raises(ValueError, match="no soportada"):
        parse_file(unsupported)
    with pytest.raises(ValueError, match="extraer texto"):
        parse_file(empty)


def test_chunk_text_validates_overlap():
    with pytest.raises(ValueError):
        chunk_text("uno dos tres", chunk_size=10, overlap=10)
